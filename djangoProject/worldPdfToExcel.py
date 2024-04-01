import pdfplumber
from docx import Document
from openpyxl import Workbook
import os
import re

# pdf_dir = 'path_to_your_pdf_directory'  # 替换为你的PDF文件所在的目录
# PDF文件路径
pdf_dir = 'C:\\PycharmProjects\\pythonProject\\worldPdf'

# Word文档路径
word_dir = 'C:\\PycharmProjects\\pythonProject\\worldPdf'

output_excel_path = 'C:\\PycharmProjects\\pythonProject\\output.xlsx'      # 输出Excel文件的路径

# 初始化Excel工作簿和工作表
wb = Workbook()
ws = wb.active
ws.append(["文件名", "关键字后一行的内容"])  # 添加表头

# 关键字
search_keyword = 'IBM, the IBM logo'

# 遍历目录中的所有PDF文件
for filename in os.listdir(pdf_dir):
    if filename.endswith('.pdf'):
        file_path = os.path.join(pdf_dir, filename)

        # 使用pdfplumber打开PDF文件
        with pdfplumber.open(file_path) as pdf:
            # 遍历PDF中的每一页
            for page in pdf.pages:
                # 提取当前页面的文本
                text = page.extract_text()

                # 搜索关键字在文本中的位置
                keyword_positions = [m.start() for m in re.finditer(search_keyword, text)]

                # 遍历所有找到的关键字位置
                for pos in keyword_positions:
                    # 尝试找到关键字所在行的结束位置（即下一个换行符的位置）
                    end_of_line = text.find('\n', pos)
                    if end_of_line == -1:  # 如果没有找到换行符，则使用页面文本的末尾
                        end_of_line = len(text)

                        # 提取关键字所在行的文本
                    current_line = text[pos:end_of_line]
                    print(f"找到关键字所在行: {current_line}")

                    # 提取后一行的文本
                    start_of_next_line = end_of_line
                    if start_of_next_line < len(text):
                        next_newline_pos = text.find('\n', start_of_next_line + 1)
                        if next_newline_pos == -1:  # 如果没有找到下一个换行符，则使用页面文本的末尾
                            next_newline_pos = len(text)
                        next_line = text[start_of_next_line + 1:next_newline_pos]
                        print(f"关键字后一行的文本: {next_line}")
                    else:
                        print("已经是页面末尾，没有后一行。")

                # # 打印提取的文本或进行其他处理
                # print(f"文件: {filename}, 页面: {page.page_number}, 文本内容:\n{text}\n")

                # 如果需要保存文本到文件，可以使用以下代码
                # with open(f'{filename}_text.txt', 'a') as text_file:
                #     text_file.write(text + '\n')

# # 注意：这个脚本将直接打印每个页面的文本内容。如果你想要保存文本内容到文件，
# # 可以取消注释保存文本到文件的代码块，并根据需要调整文件名和保存路径。

print("----------------------------------------------------------------------------")

search_keyword = '大连'

# 遍历目录下的所有Word文件
for filename in os.listdir(word_dir):
    if filename.endswith('.docx'):
        file_path = os.path.join(word_dir, filename)

        # 加载Word文档
        doc = Document(file_path)

        # 标志变量，用于跟踪是否找到了关键字
        found_keyword = False

        # 遍历文档中的每一个段落
        for para_index, para in enumerate(doc.paragraphs):
            # 检查段落中是否包含关键字
            if search_keyword in para.text:
                found_keyword = True
                # 尝试获取后一个段落的文本（确保索引不会越界）
                if para_index + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[para_index + 1]
                    print(f"在文件 {filename} 中找到关键字后的段落: {next_para.text}")
                else:
                    print(f"在文件 {filename} 中找到关键字，但已经是最后一个段落了。")
                    # 找到关键字后，可以选择继续搜索或跳出循环，这取决于你的需求
                # break  # 如果你只想找到第一个匹配就停止搜索，可以取消注释这行代码

        # 如果没有找到关键字，则打印消息
        if not found_keyword:
            print(f"在文件 {filename} 中未找到关键字。")

print("----------------------------------------------------------------------------")
# 关键字
search_keyword = '大连'

# 遍历目录下的所有Word文件
for filename in os.listdir(word_dir):
    if filename.endswith('.docx'):
        file_path = os.path.join(word_dir, filename)

        # 加载Word文档
        doc = Document(file_path)

        # 标志变量，用于记录是否找到了关键字
        found_keyword = False

        # 遍历文档中的每一个表格
        for table in doc.tables:
            for row_index, row in enumerate(table.rows):
                # 初始化列索引变量
                col_index = 0
                for cell in row.cells:
                    # 获取单元格中的文本
                    cell_text = cell.text

                    # 检查单元格文本中是否包含关键字
                    if search_keyword in cell_text:
                        found_keyword = True
                        print(f"在文件 {filename} 的表格中找到关键字在单元格中: {cell_text}")
                        # 将结果写入Excel
                        ws.append([f'{filename}', f'{cell_text}'])

                        # 获取关键字所在行的后一个单元格（如果存在）
                        if row_index + 1 < len(table.rows):
                            next_row = table.rows[row_index + 1]
                            # 确保下一行有足够的列
                            if col_index < len(next_row.cells):
                                # 获取下一行中相同列的单元格文本
                                next_cell_text = next_row.cells[col_index].text
                                print(f"关键字所在行的后一个单元格内容: {next_cell_text}")

                                # 可以在这里添加逻辑来处理找到关键字后的操作
                        break  # 如果你只想找到第一个匹配就停止搜索当前行的单元格，可以取消注释这行代码

                    # 更新列索引
                    col_index += 1

                    # 如果没有找到关键字，则打印消息
        if not found_keyword:
            print(f"在文件 {filename} 的表格中未找到关键字。")
            ws.append([f'{filename}', "关键字未找到"])

# 保存Excel文件
wb.save(output_excel_path)
print(f"Excel文件已保存至：{output_excel_path}")
